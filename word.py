


'''替换word文档图片'''
from docx import Document
from docx.shared import Inches


def center_insert_img(doc, img):
    """插入图片"""
    for paragraph in doc.paragraphs:
        # 根据文档中的占位符定位图片插入的位置
        if '<<img1>>' in paragraph.text:
            # 把占位符去掉
            paragraph.text = paragraph.text.replace('<<img1>>', '')
            run = paragraph.add_run('')
            run.add_break()
            # 添加图片并指定大小
            run.add_picture(img, width=Inches(6.2))


def save_img_to_doc(img):
    """把图片保存到doc文件中的指定位置"""
    tpl_doc = 'reports/template.docx'
    res_doc = 'reports/res/2022-03-11.docx'
    # 打开模板文件
    document = Document(tpl_doc)
    # 插入图片居中
    center_insert_img(document, img)
    # 保存结果文件
    document.save(res_doc)


def main():
    """主函数"""
    img = 'imgs/chart.png'
    save_img_to_doc(img)


if __name__ == '__main__':
    main()
